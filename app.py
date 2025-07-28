
import streamlit as st
import pdfplumber
from docx import Document
import tempfile
import os

def convert_pdf_to_word(uploaded_pdf):
    doc = Document()

    with pdfplumber.open(uploaded_pdf) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            else:
                doc.add_paragraph(f"[Page {page_num+1} has no extractable text]")

    # Save to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# ---------------- Streamlit UI ----------------

st.set_page_config(page_title="PDF to Word Converter", layout="centered")
st.title("📄 PDF to Word Converter")
st.markdown("Upload your PDF file and download it as a Word (.docx) file.")

uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

if uploaded_file is not None:
    if st.button("🔄 Convert Now"):
        with st.spinner("Converting..."):
            output_path = convert_pdf_to_word(uploaded_file)

        with open(output_path, "rb") as file:
            st.success("✅ Conversion Successful!")
            st.download_button(
                label="📥 Download Word File",
                data=file,
                file_name="converted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        os.remove(output_path)  # Clean up temporary file
